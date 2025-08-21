import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Controle Financeiro", page_icon="💰", layout="wide")

st.title("💰 Controle Financeiro Dinâmico")
st.markdown("Gerencie suas **entradas** e **saídas** de forma simples, com gráficos e exportação de relatórios.")

# ======================
# INICIALIZAR DADOS
# ======================
if "dados" not in st.session_state:
    st.session_state.dados = pd.DataFrame(columns=["Data", "Descrição", "Tipo", "Valor"])

# ======================
# FORMULÁRIO DE LANÇAMENTOS
# ======================
st.sidebar.header("➕ Novo Lançamento")
with st.sidebar.form("form_lancamento"):
    data = st.date_input("Data")
    descricao = st.text_input("Descrição")
    tipo = st.selectbox("Tipo", ["Entrada", "Saída"])
    valor = st.number_input("Valor", min_value=0.0, format="%.2f")
    adicionar = st.form_submit_button("Adicionar")

if adicionar:
    novo = pd.DataFrame([[data, descricao, tipo, valor]], columns=["Data", "Descrição", "Tipo", "Valor"])
    st.session_state.dados = pd.concat([st.session_state.dados, novo], ignore_index=True)
    st.success("✅ Lançamento adicionado com sucesso!")

# ======================
# FORMATAÇÃO DE DADOS
# ======================
df = st.session_state.dados.copy()
df["Data"] = pd.to_datetime(df["Data"], dayfirst=True)
df["Mês"] = df["Data"].dt.strftime("%B").str.capitalize()
df["Ano"] = df["Data"].dt.year
df["Valor"] = pd.to_numeric(df["Valor"], errors="coerce").fillna(0)

# ======================
# TABELA DE LANÇAMENTOS
# ======================
st.subheader("📌 Lançamentos")
st.dataframe(df.style.apply(lambda x: ["color: green" if v == "Entrada" else "color: red" for v in x["Tipo"]], axis=1))

# ======================
# GRÁFICO MENSAL
# ======================
st.subheader("📊 Resumo Mensal")
df_grafico = df.groupby(["Ano", "Mês", "Tipo"])["Valor"].sum().unstack().fillna(0)

fig, ax = plt.subplots(figsize=(10, 5))
cores = {"Entrada": "green", "Saída": "red"}
df_grafico.plot(kind="bar", ax=ax, color=[cores.get(c, "gray") for c in df_grafico.columns])
plt.title("Entradas e Saídas por Mês")
plt.ylabel("Valor (R$)")
plt.xlabel("Ano / Mês")
plt.xticks(rotation=45)
plt.tight_layout()
st.pyplot(fig)

# ======================
# EXPORTAÇÃO DE RELATÓRIOS
# ======================
st.subheader("📂 Exportar Relatório")

# Exportar para Excel
def exportar_excel():
    output = BytesIO()
    df.to_excel(output, index=False, sheet_name="Controle Financeiro")
    return output.getvalue()

st.download_button(
    "📄 Baixar Excel",
    data=exportar_excel(),
    file_name="controle_financeiro.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# Exportar para Word
def exportar_word():
    doc = Document()
    doc.add_heading("Relatório Financeiro", 0)
    doc.add_paragraph("Resumo de lançamentos:")
    tabela = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tabela.rows[0].cells
    for i, coluna in enumerate(df.columns):
        hdr_cells[i].text = coluna
    for index, row in df.iterrows():
        row_cells = tabela.add_row().cells
        for i, coluna in enumerate(df.columns):
            row_cells[i].text = str(row[coluna])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.download_button(
    "📝 Baixar Word",
    data=exportar_word(),
    file_name="relatorio_financeiro.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Exportar para PDF
def exportar_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, "Relatório Financeiro", ln=True, align="C")
    pdf.set_font("Arial", size=12)
    for index, row in df.iterrows():
        linha = f"{row['Data'].strftime('%d/%m/%Y')} - {row['Descrição']} - {row['Tipo']} - R$ {row['Valor']:.2f}"
        pdf.multi_cell(0, 10, linha)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

st.download_button(
    "📑 Baixar PDF",
    data=exportar_pdf(),
    file_name="relatorio_financeiro.pdf",
    mime="application/pdf"
)
